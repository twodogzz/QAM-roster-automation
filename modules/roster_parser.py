"""Parse QAM roster DOCX files into structured event objects."""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date, datetime, time
import logging
from pathlib import Path
import re
from typing import Iterable

from docx import Document


MONTH_MAP = {
    "JANUARY": 1,
    "JAN": 1,
    "FEBRUARY": 2,
    "FEB": 2,
    "MARCH": 3,
    "MAR": 3,
    "APRIL": 4,
    "APR": 4,
    "MAY": 5,
    "JUNE": 6,
    "JUN": 6,
    "JULY": 7,
    "JUL": 7,
    "AUGUST": 8,
    "AUG": 8,
    "SEPTEMBER": 9,
    "SEP": 9,
    "SEPT": 9,
    "OCTOBER": 10,
    "OCT": 10,
    "NOVEMBER": 11,
    "NOV": 11,
    "DECEMBER": 12,
    "DEC": 12,
}
MONTH_SHORT = ["Jan", "Feb", "Mar", "Apr", "May", "Jun", "Jul", "Aug", "Sep", "Oct", "Nov", "Dec"]

LOGGER = logging.getLogger(__name__)
MONTH_YEAR_PATTERN = re.compile(r"\b([A-Za-z]+)\s+(\d{4})\b")
VERSION_PATTERN = re.compile(r"\bv(\d+)\b", re.IGNORECASE)
DAY_ENTRY_PATTERN = re.compile(r"^\s*(\d{1,2})\.\s*")


@dataclass(frozen=True)
class RosterEvent:
    event_date: date
    day: int
    workers_raw: str
    summary: str
    description: str
    location: str
    start_time: time
    end_time: time


@dataclass(frozen=True)
class RosterDayAssignment:
    event_date: date
    workers_raw: str


@dataclass(frozen=True)
class RosterParseResult:
    source_file: Path
    year: int
    month: int
    version: str
    summary_prefix: str
    roster_events: list[RosterEvent]
    all_day_workers: list[RosterDayAssignment]
    covered_months: list[tuple[int, int]]


def extract_version_from_filename(path: Path) -> str | None:
    match = VERSION_PATTERN.search(path.name)
    return match.group(1) if match else None


def extract_version_from_doc(doc: Document) -> str | None:  # pyright: ignore[reportGeneralTypeIssues]
    for para in doc.paragraphs:
        match = VERSION_PATTERN.search(para.text)
        if match:
            return match.group(1)
    for text in _iter_table_cell_text(doc):
        match = VERSION_PATTERN.search(text)
        if match:
            return match.group(1)
    return None


def extract_roster_month_year(doc: Document) -> tuple[int, int]:  # pyright: ignore[reportGeneralTypeIssues]
    for para in doc.paragraphs:
        detected = _extract_month_year_from_text(para.text.strip())
        if detected:
            return detected
    for text in _iter_table_cell_text(doc):
        detected = _extract_month_year_from_text(text.strip())
        if detected:
            return detected
    raise ValueError("Unable to detect roster month/year from roster document.")


def parse_roster_docx(
    docx_path: str | Path,
    *,
    volunteer_name: str = "Wayne Freestun",
    event_title: str = "Wayne volunteer Queensland Air Museum",
    location: str = "Queensland Air Museum",
    shift_start: time = time(9, 0),
    shift_end: time = time(16, 30),
) -> RosterParseResult:
    path = Path(docx_path).expanduser().resolve()
    if not path.exists():
        raise FileNotFoundError(f"Roster DOCX not found: {path}")

    doc = Document(str(path))
    roster_month, roster_year = extract_roster_month_year(doc)
    version = extract_version_from_filename(path) or extract_version_from_doc(doc) or "0"
    roster_label = MONTH_SHORT[roster_month - 1]
    summary_with_version = f"{event_title} - {roster_label} v{version}"
    summary_prefix = f"{event_title} - {roster_label}"

    LOGGER.info(
        "Parsing roster DOCX '%s' for roster period %02d/%04d version v%s",
        path.name,
        roster_month,
        roster_year,
        version,
    )

    all_workers_map: dict[date, str] = {}
    volunteer_events: dict[date, str] = {}

    for text in _iter_table_cell_text(doc):
        normalized = normalize_whitespace(text)
        if not normalized:
            continue

        parsed_entry = parse_day_entry(normalized, roster_month=roster_month, roster_year=roster_year)
        if parsed_entry is None:
            continue

        event_date, workers_raw = parsed_entry
        all_workers_map[event_date] = workers_raw
        if volunteer_name.lower() in normalized.lower():
            volunteer_events[event_date] = workers_raw

    if not all_workers_map:
        raise ValueError("No roster day entries were found in the DOCX table.")

    covered_months = sorted({(entry_date.year, entry_date.month) for entry_date in all_workers_map})
    LOGGER.info(
        "Parsed %d roster day entries across %s",
        len(all_workers_map),
        ", ".join(f"{month:02d}/{year}" for year, month in covered_months),
    )
    LOGGER.info("Matched %d roster events for volunteer '%s'", len(volunteer_events), volunteer_name)

    roster_events = [
        RosterEvent(
            event_date=event_date,
            day=event_date.day,
            workers_raw=workers,
            summary=summary_with_version,
            description=f"Workers: {workers} (Roster v{version})",
            location=location,
            start_time=shift_start,
            end_time=shift_end,
        )
        for event_date, workers in sorted(volunteer_events.items())
    ]

    return RosterParseResult(
        source_file=path,
        year=roster_year,
        month=roster_month,
        version=version,
        summary_prefix=summary_prefix,
        roster_events=roster_events,
        all_day_workers=[
            RosterDayAssignment(event_date=event_date, workers_raw=workers)
            for event_date, workers in sorted(all_workers_map.items())
        ],
        covered_months=covered_months,
    )


def to_google_event_payload(event: RosterEvent, *, timezone: str) -> dict:
    start_dt = datetime.combine(event.event_date, event.start_time)
    end_dt = datetime.combine(event.event_date, event.end_time)
    return {
        "summary": event.summary,
        "location": event.location,
        "description": event.description,
        "start": {"dateTime": start_dt.isoformat(), "timeZone": timezone},
        "end": {"dateTime": end_dt.isoformat(), "timeZone": timezone},
    }


def parse_day_number(text: str) -> int | None:
    match = DAY_ENTRY_PATTERN.match(text)
    return int(match.group(1)) if match else None


def parse_day_entry(text: str, *, roster_month: int, roster_year: int) -> tuple[date, str] | None:
    match = DAY_ENTRY_PATTERN.match(text)
    if not match:
        return None

    day = int(match.group(1))
    remainder = text[match.end() :].strip()
    explicit_month_name, workers_raw = _split_explicit_month(remainder)
    actual_month, actual_year = infer_entry_month_year(
        explicit_month_name,
        roster_month=roster_month,
        roster_year=roster_year,
    )

    try:
        event_date = date(actual_year, actual_month, day)
    except ValueError as exc:
        raise ValueError(
            f"Invalid roster date '{day} {MONTH_SHORT[actual_month - 1]} {actual_year}' in cell text: {text}"
        ) from exc

    return event_date, workers_raw


def infer_entry_month_year(
    explicit_month_name: str | None,
    *,
    roster_month: int,
    roster_year: int,
) -> tuple[int, int]:
    if not explicit_month_name:
        return roster_month, roster_year

    explicit_month = MONTH_MAP.get(explicit_month_name.upper())
    if explicit_month is None:
        raise ValueError(f"Unrecognized month name in roster cell: {explicit_month_name}")

    # The roster may include only the previous, current, or next month.
    # We keep that rule explicit so malformed layouts fail loudly instead
    # of silently assigning the wrong year.
    if explicit_month == roster_month:
        return roster_month, roster_year

    previous_month = 12 if roster_month == 1 else roster_month - 1
    next_month = 1 if roster_month == 12 else roster_month + 1

    if explicit_month == previous_month:
        return explicit_month, roster_year - 1 if roster_month == 1 else roster_year
    if explicit_month == next_month:
        return explicit_month, roster_year + 1 if roster_month == 12 else roster_year

    raise ValueError(
        "Roster cell month must be the roster month, previous month, or next month. "
        f"Found '{explicit_month_name}' in a {MONTH_SHORT[roster_month - 1]} {roster_year} roster."
    )


def normalize_whitespace(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def _split_explicit_month(text: str) -> tuple[str | None, str]:
    if not text:
        return None, ""

    first_word, separator, remainder = text.partition(" ")
    if first_word.upper() in MONTH_MAP:
        return first_word, remainder.strip()
    return None, text


def _extract_month_year_from_text(text: str) -> tuple[int, int] | None:
    match = MONTH_YEAR_PATTERN.search(text)
    if not match:
        return None
    month = MONTH_MAP.get(match.group(1).upper())
    if month is None:
        return None
    return month, int(match.group(2))


def _iter_table_cell_text(doc: Document) -> Iterable[str]:  # pyright: ignore[reportGeneralTypeIssues]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield cell.text
